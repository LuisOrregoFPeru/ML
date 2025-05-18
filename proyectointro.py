import streamlit as st
import openai
import requests
from docx import Document
from datetime import date
from typing import List

# ---------- Helper functions ---------- #

def init_openai(api_key: str):
    """Initialize OpenAI with user-provided key or stop execution."""
    if api_key:
        openai.api_key = api_key
    else:
        st.error("Por favor ingresa una clave de API de OpenAI válida.")
        st.stop()


def generate_introduction(title: str, objective: str, word_target: int = 4500) -> str:
    """Generate the thesis introduction (~9 A4 pages ≈ 4 500 words)."""

    prompt = (
        "Eres un asistente experto en redacción académica. "
        "Redactarás una introducción de proyecto de tesis basándote en el siguiente título y objetivo general. "
        "La introducción se construirá en prosa, sin subtítulos, en tercera persona y tiempo futuro indicativo. "
        "Cada párrafo tendrá un máximo de 10 líneas y el texto completo tendrá al menos 9 páginas A4 (~4500+ palabras). "
        "La estructura de la introducción será:\n"
        "1. Importancia y plausibilidad del problema (1-2 párrafos).\n"
        "2. Impacto global, regional y nacional (1-2 párrafos).\n"
        "3. Vacíos de la literatura (1-2 párrafos).\n"
        "4. Contribución a los ODS (1 párrafo).\n"
        "5. Pregunta de investigación (1 párrafo, forma interrogativa).\n"
        "6. Justificación teórica (1 párrafo).\n"
        "7. Justificación práctica (1 párrafo).\n"
        "8. Justificación metodológica (1 párrafo).\n"
        "9. Justificación social (1 párrafo).\n"
        "Después de la introducción debes escribir '===ANTECEDENTES===' y esperar a que el sistema agregue antecedentes.\n"
        f"Título: {title}\n"
        f"Objetivo general: {objective}\n"
        f"Longitud mínima: {word_target} palabras.\n"
        "Devuelve solo el texto de la introducción."
    )

    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=8192  # Ajusta si tu plan admite menos tokens
    )
    return response.choices[0].message["content"]


def search_pubmed(query: str, n: int = 10) -> List[dict]:
    """Retrieve PubMed abstracts via NCBI E-utilities."""

    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    ids = requests.get(
        base + "esearch.fcgi",
        params={"db": "pubmed", "term": query, "retmax": n, "retmode": "json"}
    ).json()["esearchresult"]["idlist"]

    abstracts = []
    for pmid in ids:
        xml = requests.get(base + "efetch.fcgi", params={"db": "pubmed", "id": pmid, "retmode": "xml"}).text
        import re, html
        title = re.search(r"<ArticleTitle>(.*?)</ArticleTitle>", xml, re.S)
        abstract = re.search(r"<AbstractText.*?>(.*?)</AbstractText>", xml, re.S)
        authors = re.findall(r"<LastName>(.*?)</LastName>.*?<Initials>(.*?)</Initials>", xml, re.S)
        year = re.search(r"<PubDate>.*?<Year>(\\d{4})</Year>", xml, re.S)
        if not (title and abstract):
            continue
        if authors:
            first = f"{authors[0][0]}, {authors[0][1][0]}."
            author_cite = first + (" et al." if len(authors) > 3 else "")
        else:
            author_cite = "Autor desconocido"
        pub_year = year.group(1) if year else "s.f."
        abstracts.append({
            "cite": f"{author_cite} ({pub_year})",
            "abstract": html.unescape(abstract.group(1))
        })
    return abstracts[:n]


def paraphrase(text: str, api_key: str, word_count: int = 130) -> str:
    """Paraphrase a PubMed abstract to ~130 words in Spanish."""

    prompt = (
        f"Parafrasea el siguiente resumen en español académico, sin plagio, en aproximadamente {word_count} palabras:\n{text}"
    )
    response = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=512
    )
    return response.choices[0].message["content"]


def build_antecedents(abstracts: List[dict], api_key: str) -> str:
    """Compose antecedent paragraphs from abstracts."""

    paragraphs = []
    for item in abstracts:
        para = paraphrase(item["abstract"], api_key)
        paragraphs.append(f"{item['cite']} {para}")
    return "\n\n".join(paragraphs)


def generate_theoretical_bases(title: str, objective: str, api_key: str) -> str:
    """Generate theoretical framework section."""

    prompt = (
        "Redacta las bases teóricas pertinentes a la pregunta de investigación, "
        "incluyendo enfoques conceptuales y variables clave. "
        "Redacta en prosa académica, tercera persona, futuro indicativo:\n"
        f"Título: {title}\nObjetivo general: {objective}"
    )
    resp = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=2048
    )
    return resp.choices[0].message["content"]


def generate_hypotheses(objective: str, api_key: str) -> str:
    """Formulate research and statistical hypotheses."""

    prompt = (
        "Formula las hipótesis de investigación y estadísticas (nula y alternativa) "
        "basadas en el siguiente objetivo general. "
        "Presenta primero la hipótesis de investigación y luego las estadísticas separadas:\n"
        f"{objective}"
    )
    resp = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=512
    )
    return resp.choices[0].message["content"]


def build_docx(intro: str, antecedentes: str, teoricas: str, hipotesis: str) -> bytes:
    """Build a DOCX with all generated sections and return it as bytes."""

    doc = Document()
    doc.add_heading("Proyecto de Tesis – Secciones Generadas", level=1)

    doc.add_heading("Introducción", level=2)
    for p in intro.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Antecedentes", level=2)
    for p in antecedentes.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Bases Teóricas", level=2)
    for p in teoricas.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Hipótesis", level=2)
    for p in hipotesis.split("\n"):
        doc.add_paragraph(p)

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ---------- Streamlit UI ---------- #

st.set_page_config(page_title="Generador de Introducciones de Tesis", layout="wide")
st.title("📝 Generador Automático de Introducciones de Tesis")

with st.sidebar:
    st.header("Configuración")
    api_key = st.text_input("OpenAI API Key", type="password")
    title_input = st.text_input("Título de la Investigación")
    objective_input = st.text_area("Objetivo General")
    generate_btn = st.button("Generar Introducción")

if generate_btn:
    init_openai(api_key)

    with st.spinner("Generando introducción..."):
        intro_text = generate_introduction(title_input, objective_input)

    st.subheader("Introducción")
    st.markdown(intro_text)

    with st.spinner("Buscando antecedentes en PubMed..."):
        abstracts = search_pubmed(title_input or objective_input)
        antecedentes_text = build_antecedents(abstracts, api_key)

    st.subheader("Antecedentes")
    st.markdown(antecedentes_text)

    with st.spinner("Generando bases teóricas..."):
        teoricas_text = generate_theoretical_bases(title_input, objective_input, api_key)

    st.subheader("Bases Teóricas")
    st.markdown(teoricas_text)

    with st.spinner("Generando hipótesis..."):
        hipotesis_text = generate_hypotheses(objective_input, api_key)

    st.subheader("Hipótesis")
    st.markdown(hipotesis_text)

    docx_bytes = build_docx(intro_text, antecedentes_text, teoricas_text, hipotesis_text)
    st.download_button(
        label="Descargar DOCX",
        data=docx_bytes,
        file_name=f"Tesis_{date.today()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("""
---
### ¿Cómo funciona?
1. Ingresa tu clave de API de OpenAI.
2. Escribe el título o el objetivo general de tu investigación.
3. Haz clic en **Generar Introducción** y espera los resultados.
4. Descarga el documento final en formato **.docx**.
""")
