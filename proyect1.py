import streamlit as st
import requests
from docx import Document
from typing import List
import textwrap

# ---------------- Configuración por defecto ---------------- #
# Modo simple: sin modelo LLM ni tokens API

# ---------------- Cliente Hugging Face (desactivado) ---------------- #

def get_client(token=None, model_id=None):
    """Siempre retorna None para forzar el modo simple."""
    return None


def hf_generate(client, prompt: str, max_tokens: int = 1024, temperature: float = 0.3) -> str:
    """Wrapper que siempre llama al generador simple."""
    return simple_llm(prompt)

# ---------------- Modo simple ---------------- #

def simple_paragraph(text: str) -> str:
    """Ajusta el texto para no exceder ~100 caracteres por línea."""
    return textwrap.fill(text, width=100)


def simple_llm(prompt: str) -> str:
    """Genera texto placeholder cuando no hay modelo disponible."""
    placeholders = [
        "Se resaltará la trascendencia del fenómeno planteado y la pertinencia científica de abordarlo.",
        "Se describirá la magnitud del problema a escala global, regional y nacional, destacando sus repercusiones sanitarias y económicas.",
        "Se sintetizarán las principales brechas de conocimiento detectadas en la literatura, subrayando la necesidad de nuevos estudios.",
        "La investigación contribuirá al Objetivo de Desarrollo Sostenible relacionado con la salud y el bienestar.",
        "¿En qué medida el fenómeno descrito se asociará con las variables seleccionadas en el contexto propuesto?",
        "Desde el punto de vista teórico, el estudio profundizará en los modelos conceptuales vigentes y propondrá nuevas perspectivas.",
        "En el plano práctico, los hallazgos orientarán intervenciones basadas en evidencia para mejorar la situación analizada.",
        "Metodológicamente, se empleará un diseño robusto que garantizará la validez y confiabilidad de los resultados obtenidos.",
        "El beneficio social radicará en la generación de conocimiento aplicable que favorecerá la calidad de vida de la población implicada."
    ]
    return "\n\n".join(simple_paragraph(p) for p in placeholders)

# ---------------- Funciones principales ---------------- #

def generate_introduction(client, title: str, objective: str, word_target: int = 4500) -> str:
    prompt = f"""
Redacta una introducción de tesis académica con las siguientes características:
• Extensión mínima: {word_target} palabras (≈ 9 páginas A4).
• Prosa, sin subtítulos, tercera persona, tiempo futuro del modo indicativo.
• Máximo 10 líneas por párrafo.
• Secuencia de párrafos exacta:
  1‑2 p Importancia + plausibilidad del problema.
  1‑2 p Impacto (mundo, Latinoamérica, Perú).
  1‑2 p Vacíos de literatura.
  1 p Contribución a ODS.
  1 p Pregunta de investigación (interrogativa).
  1 p Justificación teórica.
  1 p Justificación práctica.
  1 p Justificación metodológica.
  1 p Justificación social.
Al final escribe la línea EXACTA "===ANTECEDENTES===".
Título: {title}
Objetivo general: {objective}
"""
    return hf_generate(client, prompt)


def paraphrase_abstract(client, abstract: str, word_count: int = 130) -> str:
    """Devuelve texto genérico en modo simple."""
    return simple_paragraph("Resumen disponible para consulta; se presentará de forma sintética y libre de plagio en la versión final.")


def search_pubmed(query: str, n: int = 10) -> List[dict]:
    """Obtiene n resúmenes de PubMed (sin token)."""
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    ids = requests.get(base + "esearch.fcgi", params={"db": "pubmed", "term": query, "retmax": n, "retmode": "json"}).json()["esearchresult"]["idlist"]
    items = []
    for pmid in ids:
        xml = requests.get(base + "efetch.fcgi", params={"db": "pubmed", "id": pmid, "retmode": "xml"}).text
        import re, html
        abstract = re.search(r"<AbstractText.*?>(.*?)</AbstractText>", xml, re.S)
        if not abstract:
            continue
        authors = re.findall(r"<LastName>(.*?)</LastName>.*?<Initials>(.*?)</Initials>", xml, re.S)
        year = re.search(r"<PubDate>.*?<Year>(\\d{4})</Year>", xml, re.S)
        if authors:
            first = f"{authors[0][0]}, {authors[0][1][0]}."
            cite = first + (" et al." if len(authors) > 3 else "")
        else:
            cite = "Autor desconocido"
        items.append({
            "cite": f"{cite} ({year.group(1) if year else 's.f.'})",
            "abstract": html.unescape(abstract.group(1))
        })
    return items[:n]


def build_antecedents(client, pubs: List[dict]) -> str:
    return "\n\n".join(
        f"{p['cite']} {paraphrase_abstract(client, p['abstract'])}" for p in pubs
    ) or "Sin referencias disponibles."


def generate_theoretical_bases(client, title: str, objective: str) -> str:
    return simple_paragraph("Se desarrollarán los fundamentos conceptuales que sustentan la relación entre las variables propuestas y se explicará el marco teórico que guiará el análisis.")


def generate_hypotheses(client, objective: str) -> str:
    return simple_paragraph("Hipótesis de investigación y estadísticas serán formuladas relacionando las variables primarias para demostrar la dirección y fuerza del efecto esperado.")


def build_docx(intro: str, antecedentes: str, bases: str, hyps: str) -> bytes:
    doc = Document()
    doc.add_heading("Proyecto de Tesis – Secciones Generadas", 1)
    doc.add_heading("Introducción", 2)
    for p in intro.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Antecedentes", 2)
    for p in antecedentes.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Bases Teóricas", 2)
    for p in bases.split("\n"):
        doc.add_paragraph(p)
    doc.add_page_break()

    doc.add_heading("Hipótesis", 2)
    for p in hyps.split("\n"):
        doc.add_paragraph(p)

    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# ---------------- Interfaz Streamlit ---------------- #

st.set_page_config(page_title="Generador de Introducciones de Tesis", layout="wide")
st.title("📝 Generador Automático de Introducciones de Tesis – Solo Modo Simple")

with st.sidebar:
    st.header("Configuración")
    title_input = st.text_input("Título de la Investigación")
    objective_input = st.text_area("Objetivo General")
    generate_btn = st.button("Generar Introducción")

if generate_btn:
    client = get_client()  # Siempre None

    with st.spinner("Generando introducción…"):
        intro = generate_introduction(client, title_input, objective_input)
    st.subheader("Introducción")
    st.markdown(intro)

    with st.spinner("Buscando antecedentes en PubMed…"):
        pubs = search_pubmed(title_input or objective_input) if (title_input or objective_input) else []
        antecedentes = build_antecedents(client, pubs)
    st.subheader("Antecedentes")
    st.markdown(antecedentes)

    with st.spinner("Generando bases teóricas…"):
        bases = generate_theoretical_bases(client, title_input, objective_input)
    st.subheader("Bases Teóricas")
    st.markdown(bases)

    with st.spinner("Generando hipótesis…"):
        hyps = generate_hypotheses(client, objective
